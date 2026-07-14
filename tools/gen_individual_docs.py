import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = "Microsoft YaHei"
BASE = "h:/CampusRPG_作业版/"
ASSETS = BASE + "assets/"

def F(run, size=12, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        s = {0:26,1:18,2:15,3:13}.get(level,13)
        F(r, size=s, bold=True)

def P(doc, text, size=12, bold=False, indent=0, color=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        F(run, size=size, bold=bold, color=color)

def IMG(doc, path, width=4.5, caption=""):
    """Insert image with optional caption, centered."""
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    except:
        return
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        F(r, size=9, color=(100,100,100))

def IMGS(doc, paths, width=1.8, per_row=3):
    """Insert multiple images in a grid."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    count = 0
    for path in paths:
        if os.path.exists(path):
            try:
                run = p.add_run()
                run.add_picture(path, width=Inches(width))
                run = p.add_run("  ")
                count += 1
                if count % per_row == 0:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass

def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: F(r, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: F(r, size=9)
    return tbl

def title_page(doc, name, role, pct):
    for _ in range(5): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("校园RPG冒险游戏"); F(r, size=36, bold=True)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(f"C++课程设计答辩"); F(r, size=20, bold=True, color=(80,80,80))
    doc.add_paragraph()
    n = doc.add_paragraph(); n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = n.add_run(f"答辩人：{name}（{role}）"); F(r, size=16, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"贡献度：{pct}"); F(r, size=14, color=(100,100,100))
    for _ in range(2): doc.add_paragraph()
    b = doc.add_paragraph(); b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = b.add_run("计科2503班  -  2026年7月"); F(r, size=14, color=(100,100,100))
    doc.add_page_break()

# =====================================================
members = [
    {
        "name": "谢文浩",
        "role": "组长",
        "pct": "30%",
        "modules": "系统架构设计 + 回合制战斗引擎 + 地窟生成系统 + 等级成长 + Git仓库管理",

        "work_table": [
            ["系统架构设计", "定义MainWindow.h全部284行数据结构（CharacterData/EnemyData/ItemData/TaskData/RoomData）；设计模块间接口规范；设计Profession抽象基类到6个派生职业的继承体系；选择QVector+QMap作为主要容器"],
            ["回合制战斗引擎", "实现startBattle()/fightOneRound()/enemyTurn()/endBattleIfNeeded()完整状态机；先手判定算法（双方队伍平均等级比较）；伤害公式max(1,攻击-防御)；前后排目标智能选择（敌人优先攻击前排）；6种异常状态系统（冻结跳过行动/灼烧持续掉血/减速降低速度/眩晕无法行动/沉默禁用技能/嘲讽强制攻击）"],
            ["地窟生成系统", "buildDungeonLayer()根据层数生成房间；7种房间类型（恶魔商店/普通战斗x3/宝箱/精英通道/BOSS房）；makeEnemyGroup()按7层x3类(普通/精英/BOSS)生成30+种敌人，每层敌人属性递增；makeLayerEquipments()按层生成4槽位装备掉落表"],
            ["等级成长系统", "addExp()累计经验、检测升级；属性随等级按职业系数自动增长（不同职业成长率不同）；技能按Lv1/5/10/15/20五档自动解锁；expNeed随等级递增（100→150→225→...）"],
            ["Git仓库管理", "制定分支策略（main+feature分支）；多成员代码合并与冲突解决；GitHub Release版本打包发布；处理多次force push解决历史分叉问题"],
        ],

        "speech": """各位老师好，我是谢文浩，本项目的组长，主要负责系统架构设计和核心游戏逻辑的开发。

首先介绍系统架构。整个项目的核心数据结构定义在MainWindow.h中，我设计了CharacterData、EnemyData、ItemData、TaskData、RoomData等核心结构体，以及GamePhase、RoomType等枚举类型。选择struct而非class是因为游戏数据需要频繁在模块间传递，public访问更便捷。同时设计了Profession抽象基类的继承体系——6个职业继承同一个基类，各自重写skills()和baseStats()虚函数，实现运行时多态。

下面重点介绍战斗系统，这是整个项目技术难度最高的部分。战斗流程是一个完整的状态机：startBattle初始化战斗→fightOneRound我方行动→enemyTurn敌方行动→endBattleIfNeeded判断胜负。先手判定算法是取双方队伍平均等级比较。伤害公式采用经典的max(1,攻击-防御)保证至少1点伤害。

战斗系统最核心的创新是6种异常状态机制。冻结会让目标跳过行动、灼烧每回合掉血、减速降低出手速度、眩晕无法行动、沉默禁止使用技能、嘲讽强制攻击特定目标。这些状态通过EnemyData中的frozen/slow/burn等字段追踪持续回合数，在enemyTurn中统一处理。

地窟系统我生成了7层共40+个房间。每层调用makeEnemyGroup()按普通/精英/BOSS三类生成不同难度的敌人，属性和技能随层数递增。比如第一层敌人HP约100，到第七层BOSS达到了1200点，保证难度曲线的平滑上升。

以上就是我的答辩内容，谢谢各位老师。""",

        "qa": [
            ("为什么选择Qt Widgets而不是Qt Quick/QML？",
             "Qt Widgets是C++课程设计的主流选择，与课程教学内容（C++面向对象编程）结合更紧密。Widgets的所有控件都是C++类，可以直接体现封装、继承、多态等OOP特性。而QML需要学习JavaScript/QML语法，偏离了C++课程的核心目标。"),
            ("战斗中异常状态是用什么数据结构管理的？",
             "异常状态通过EnemyData结构体中的整数字段管理，如frozen表示冻结剩余回合、burn表示灼烧剩余回合。在enemyTurn()中检查这些字段，>0则触发效果并递减。这种方式比用vector<StatusEffect>更轻量，因为每个敌人同时只有1-2个状态，用6个int比动态容器更高效。"),
            ("地窟难度递增是如何保证的？",
             "通过makeEnemyGroup()的layer参数控制。每层敌人的HP、攻击、防御按固定系数递增（约15%-25%），同时给高层敌人添加更多技能词条。装备掉落也随层数提升品质，例如第一层掉落+5属性的装备，第七层掉落+35属性的装备。"),
            ("如果新增一个职业，需要修改哪些代码？",
             "只需两步：1. 创建新类继承Profession，实现skills()和baseStats()两个虚函数；2. 在setupData()中push_back到professions列表。MainWindow的战斗、编队、成长等所有逻辑都通过基类指针操作，不需要任何修改。这就是多态的核心优势。"),
            ("项目中最难解决的技术问题是什么？",
             "战斗状态机的设计是最难的。需要处理好玩家选择目标→攻击→敌方反击→状态更新→胜负判定的完整流程，还要处理玩家或敌人在中途死亡的情况。我采用了分阶段处理：先收集所有存活角色→按速度排序→依次执行动作→每步检查死亡→回合结束统一处理异常状态。这样保证了逻辑的清晰性和正确性。"),
            ("Git合并时遇到冲突怎么处理？",
             "首先用git status查看冲突文件列表，然后逐个打开冲突文件，根据冲突标记<<< === >>>判断谁改了什么。对于MainWindow.cpp这种多人协作的大文件，我会和组员确认各自的修改范围，合并后编译验证确保没有遗漏。最后强制推送时先备份远程最新版本，避免覆盖别人的工作。"),
        ],
        "images_title": "战斗系统与地窟模型展示",
        "images_desc": "以下为负责模块相关的游戏素材：最终BOSS、精英敌人、地窟场景",
        "images": [
            "sprites/boss_principal.png",
            "sprites/elite_book_guardian.png",
            "sprites/elite_clock_prisoner.png",
            "sprites/elite_mutant.png",
            "sprites/elite_rust_regret.png",
            "sprites/elite_winged_runner.png",
            "sprites/elite_faceless_dancer.png",
            "scenes/dungeon_07_boss.png",
            "scenes/gym.jpg",
        ],
        "img_width": 1.6,
        "img_per_row": 3,
    },
    {
        "name": "黄金弘",
        "role": "组员",
        "pct": "24%",
        "modules": "天使商店系统 + 恶魔商店系统 + 任务系统 + 数据持久化（存档序列化）",

        "work_table": [
            ["天使商店", "refreshAngelShop()/buyAngelItem()实现药品/食品/养成类共14种商品；折扣系统（angelDiscount每次+1%，最高50%）；金币交易与库存检测；firstAngelShopBought标记首次购买"],
            ["恶魔商店", "refreshDemonShop()/buyDemonItem()/sellEquipmentToDemon()消耗品/装备/恶魔之友；demonCoin独立货币系统；装备半价回收机制；demonDiscount累计；demonFriendBonus加成"],
            ["任务系统", "refreshTasks()/acceptTask()/claimTask()完整生命周期管理；3种条件类型（DefeatEnemy/CollectItem/DefeatBoss）；checkTaskCompletion()自动进度检测；addTaskProgress()进度累加；奖励发放（金币+恶魔币+物品）"],
            ["数据持久化", "writeGame()/loadGame()基于INI文本的存档格式；4个存档槽位独立文件(slot_1.ini~slot_4.ini)；8个序列化/反序列化函数：serializeCharacters/deserializeCharacters、serializeInventory/deserializeInventory、serializeTasks/deserializeTasks、serializeCodex/deserializeCodex"],
        ],

        "speech": """各位老师好，我是黄金弘，负责商店系统、任务系统和数据持久化三个模块。

首先介绍商店系统。游戏中有两种商店：天使商店使用金币交易，恶魔商店在地窟中使用恶魔币交易。天使商店有14种商品，分为药品/食品/养成三类。我设计了折扣累计机制——每次在天使商店购买后，折扣+1%，最高50%，这鼓励玩家持续在同一商店消费。恶魔商店的恶魔币只能在地窟战斗中获取，装备可以半价卖回给恶魔商店，形成了独立的经济循环。

接下来是任务系统。我实现了完整的任务生命周期：未接→已接→已完成→已领奖。任务有三种条件类型：击败指定敌人、收集指定物品、击败最终BOSS。最关键的设计是自动检测机制——checkTaskCompletion()在玩家击败敌人或获得物品后自动触发，遍历所有已接任务检查条件是否满足，满足则状态变为已完成。这避免了玩家需要手动"交任务"的繁琐操作。

最后是数据持久化，这是保证游戏体验的核心。我选择了INI文本格式而非二进制格式，原因是：第一，可读性强，打开slot_1.ini就能看到角色名、等级、HP、背包物品等所有数据；第二，方便调试，老师和我们可以直接查看存档内容验证正确性；第三，格式简单，不易出错。

存档系统有4个独立槽位，每个槽位对应一个slot_X.ini文件。序列化时我设计了8个函数，分成4组：角色序列化（队伍数据）、物品序列化（背包数据）、任务序列化（任务状态和进度）、图鉴序列化（已遭遇敌人）。分组的原因是每组数据格式不同，分开处理更清晰，出问题时也容易定位。反序列化时按同样顺序读取，每个字段都做了格式校验，防止存档损坏导致崩溃。

以上就是我的答辩内容，谢谢各位老师。""",

        "qa": [
            ("为什么选择INI格式而不是JSON或XML？",
             "INI格式是Qt的QSettings原生支持的格式，读取写入都不需要引入第三方库。JSON需要QJsonDocument，XML需要QXmlStreamReader，复杂度更高。INI的section=key=value结构对于游戏存档来说足够了——游戏数据本质就是键值对的集合。而且INI文件可以直接用记事本打开编辑，方便老师和同学验证。"),
            ("8个序列化函数为什么不用一个统一函数？",
             "单一函数会导致单一职责原则被破坏。角色/物品/任务/图鉴的数据结构完全不同的，序列化格式也不同。如果合并成一个函数，需要用大量的if-else或switch来区分类型，代码会变得臃肿难维护。分开后每组函数职责明确，出问题时能快速定位到具体是哪类数据读写错误。"),
            ("任务进度自动检测会不会有性能问题？",
             "不会。checkTaskCompletion()只在两类事件后触发：背包物品变化时和击败敌人时。触发频率很低（玩家手动操作），每次只需遍历3-4个已接任务，每个任务做O(1)的条件判断。时间复杂度可以忽略不计。"),
            ("如果存档文件被手动修改导致格式错误怎么办？",
             "loadGame()的每个反序列化函数都有try-catch保护。如果某个字段解析失败（比如数字字段读到了文字），整个loadGame()返回false，提示存档损坏，不会导致崩溃。同时原始文件保留，用户可以把备份的存档覆盖回来。"),
            ("双货币系统是如何平衡的？",
             "金币和恶魔币是两个完全独立的货币体系，互不兑换。金币通过打工和任务获取，在天使商店消费；恶魔币通过地窟战斗获取，在恶魔商店消费。这样设计保证了地窟探索的必要性——玩家不能通过反复打工来购买恶魔商店的强力装备，必须深入地窟冒险。"),
            ("恶魔商店的装备半价回收逻辑是如何实现的？",
             "sellEquipmentToDemon()中，首先列出背包中的所有装备类物品，玩家选择要出售的装备编号，系统将该装备的demonPrice除以2作为回收价格，增加玩家恶魔币，然后从背包中移除该物品。折扣系统(demonDiscount)不影响回收价，回收价始终是原价的一半。"),
        ],
        "images_title": "商店系统场景展示",
        "images_desc": "以下为负责模块相关的游戏场景：天使商店、恶魔商店所在的地窟环境",
        "images": [
            "scenes/dungeon_01_gym.jpg",
            "scenes/dungeon_02_library.jpg",
            "scenes/dungeon_03_theater.jpg",
            "scenes/dungeon_04_lab.jpg",
            "scenes/dungeon_05_divination.jpg",
            "scenes/dungeon_06_office.jpg",
            "scenes/demon_shop.jpg",
        ],
        "img_width": 1.6,
        "img_per_row": 3,
    },
    {
        "name": "张轩",
        "role": "组员",
        "pct": "23%",
        "modules": "角色管理系统 + 背包系统 + 装备穿戴 + 编队阵型",

        "work_table": [
            ["角色管理", "createCharacter()6职业创建流程（选择职业→输入名称→自动赋值属性）；refreshCharacters()角色详情展示（图标/属性/技能/装备槽）；applyProfessionStats()根据职业差异化赋值基础属性（HP/MP/攻/防/魔攻/法抗/活力）"],
            ["背包系统", "refreshInventory()40格背包显示与容量检测；useInventoryItem()战斗内外使用逻辑分离（食品仅战斗外、药品无限制、装备不可使用）；discardInventoryItem()删除确认；addItem()容量溢出提示；inventoryText()物品信息格式化"],
            ["装备系统", "equipSelectedItem()穿戴/卸下切换逻辑；4种装备槽位（物攻/法攻/物防/法抗）独立管理，同槽位替换时自动归还旧装备；装备属性百分比加成计算（如装备15%物攻→实际物攻+15%）；装备状态跟踪equippedOnce标记"],
            ["编队系统", "setFormationOneFront()前排1后排2模式；setFormationTwoFront()前排2后排1模式切换；moveRoleUp()/moveRoleDown()角色上下排序调整；toggleSelectedRoleActive()角色启用/禁用；formationType记录当前阵型"],
        ],

        "speech": """各位老师好，我是张轩，负责角色管理、背包系统、装备穿戴和编队系统四个模块。

首先介绍角色管理模块。玩家在创建角色时可以选择6个职业——学生、冰法师、圣骑士、祈福者、血战士和魔术师。每个职业有完全不同的基础属性：圣骑士是肉盾型，HP高达160；血战士是近战输出型，攻击力26；魔术师是远程法师型，魔法攻击28但HP只有85。这种差异化设计让每个职业都有独特的游戏体验。创建时调用applyProfessionStats()根据职业名称自动赋值所有属性。

然后是背包系统。背包上限40格，物品分为药品、食品、养成、消耗品、装备五类。战斗中有一个重要的限制设计——食品不能在战斗中使用。这个设计是为了增加战斗的策略性：玩家必须在战斗前决定携带多少药品，不能在战斗中靠食品无限回血。useInventoryItem()接收一个inBattle参数来判断当前是否在战斗中，如果是战斗且物品类型是食品，直接拒绝使用。

装备系统支持4种槽位：物攻装备、法攻装备、物防装备、法抗装备。每个槽位只能穿戴一件装备，再次穿戴同槽位装备时会自动卸下旧装备并放回背包。装备提供的是百分比加成而非固定数值，比如"幕布匕首+15%物攻"——这保证了装备在整个游戏过程中都有价值，不会因为等级提升而贬值。

最后是编队系统。游戏支持两种阵型：前排1后排2（一个主坦在前吸收伤害）和前排2后排1（双前排分担伤害）。战斗中的敌人AI会优先攻击前排角色，所以阵型的选择直接影响战斗策略。玩家还可以用上移/下移按钮调整角色顺序，以及禁用不需要的角色。

以上就是我的答辩内容，谢谢各位老师。""",

        "qa": [
            ("为什么装备采用百分比加成而不是固定数值？",
             "百分比加成保证装备的长期价值。如果装备+10攻击力，角色等级提升后这10点就微不足道了。但+15%物攻始终能让角色的攻击力提升15%，无论角色是1级还是30级。这是RPG游戏设计的常见做法。"),
            ("食品为什么不能在战斗中使用？",
             "为了增加战斗的策略性和紧张感。如果战斗中能使用食品（通常比药品便宜且量大），玩家就可以无限续战，战斗失去挑战性。限制食品只能在战斗外使用，玩家在进入地窟前必须规划好携带多少药品，这本身就是策略的一部分。"),
            ("编队系统如何影响战斗AI？",
             "敌人的攻击逻辑是优先攻击前排角色。前排1后排2模式下，前排角色承受70%以上的攻击；前排2后排1模式下，两个前排分担伤害。enemyTurn()中通过遍历我方存活角色列表，优先选择active=true且处于前排的角色作为攻击目标。只有拥有backlineAttack能力的精英敌人才会攻击后排。"),
            ("角色属性中'活力'的作用是什么？",
             "活力是排课系统中的核心资源。每个角色初始活力100，上课/打工/补习班都会消耗活力，周末选择'不上课'可以恢复15点活力。活力不足时无法执行消耗大的行动。活力也影响战斗表现——活力越低，角色在地窟中初始HP会有一定折扣。这模拟了学生精力管理的真实场景。"),
            ("如果背包满了想捡装备怎么办？",
             "addItem()会先检查当前背包大小是否达到40格上限。如果满了，会弹出提示'背包已满'，物品不会丢失但也不会加入背包。玩家需要先使用或丢弃背包中的物品腾出空间。战斗中掉落的物品如果背包满了，也会提示无法获得，鼓励玩家定期清理背包。"),
            ("装备卸下后属性会不会计算错误？",
             "不会。装备穿戴和卸下都通过equipSelectedItem()统一处理。穿戴时记录装备名到role.equipment[slotName]，计算属性时遍历所有槽位累加百分比加成。卸下时从equipment中移除对应槽位，属性加成自动消失。角色每次刷新都重新计算总属性，不会出现累积错误。"),
        ],
        "images_title": "六职业角色模型展示",
        "images_desc": "以下为负责模块相关的6个职业角色设计模型（AI生成概念图）",
        "images": [
            "sprites/学生.jpg",
            "sprites/冰法师.jpg",
            "sprites/圣骑士.jpg",
            "sprites/祈福者.jpg",
            "sprites/血战士.jpg",
            "sprites/魔术师.jpg",
        ],
        "img_width": 1.8,
        "img_per_row": 3,
    },
    {
        "name": "张申桥",
        "role": "组员",
        "pct": "23%",
        "modules": "Qt图形界面开发 + 排课系统 + 系统集成测试 + 文档撰写与资源管理",

        "work_table": [
            ["Qt界面开发", "setupSavePage()存档首页（4个存档槽位+新建/删除按钮）；setupGamePage()游戏主页（QTabWidget+8个Tab页）；暗色主题CSS样式表（背景#1e1e2e/文字#d4d4d4/按钮#3a3a5c/选中#5a5a8c）；updateUiScale()随窗口大小自适应缩放字体和图标；QListWidget图标显示（图鉴/编队/地窟含精灵图缩略图）"],
            ["排课系统", "rebuildScheduleTable()7天x2半天=14格课表QTableWidget生成；日程下拉选择（工作日9科+周末3种）；refreshSchedulePreview()本周预计收益实时预览（活力消耗+金币收入计算）；actionVigorCost()/actionGoldChange()对应课程的效果公式；补习班规则（效果减半effectScale=0.5）"],
            ["集成测试", "全流程端到端测试：新建角色→排课→天使商店→地窟7层→BOSS战→通关/团灭轮回；13项测试用例覆盖存档读写/6职业创建/排课消耗/商店买卖/地窟探索/战斗引擎/任务检测/异常状态/装备穿戴/编队切换/图鉴记录/轮回保留/通关判定"],
            ["文档与资源", "README.md完整项目文档（功能列表/编译指南/OOP特性/STL容器说明）；检查点一/二/三课程设计报告；AI生成72张素材整理与规范命名；免安装启动包（EXE+DLL+platforms+saves）；答辩Word文档和分工文档"],
        ],

        "speech": """各位老师好，我是张申桥，负责整个项目的Qt图形界面开发、排课系统、集成测试和文档资源管理。

首先展示Qt界面架构。我采用QStackedWidget实现双页面切换——首页是存档选择页，有4个存档槽位，显示角色名或"空白存档"；点击进入后是游戏主页面，使用QTabWidget包含8个标签页：轮回总览、排课、角色编队、背包、任务、天使商店、地窟/恶魔商店和轮回游记。

界面的视觉设计我选择了暗色主题，背景色#1e1e2e，文字#d4d4d4，按钮#3a3a5c带悬停高亮。所有字体通过updateUiScale()随窗口大小自适应缩放——以基准尺寸1280x820为参考，按比例调整所有控件的字体大小，保证在不同分辨率下都有良好的阅读体验。

一个重要的UI创新是在图鉴和编队页面加入了精灵图图标显示。通过QListWidget的setIcon方法，将assets/sprites目录下的PNG精灵图加载为64-136像素的缩略图，让玩家在文字信息之外能直观看到角色和敌人的形象。

排课系统是这个项目特有的功能模块。我用QTableWidget生成了7天x2半天的14格课表，每个格子是QComboBox下拉选择框。工作日可选择9种课程（专业课/大物课/高数课等），周末只允许不上课/打工/补习班三种。选择后实时显示本周预计盈利和剩余活力，帮助学生玩家做出最优的排课决策。

测试方面，我做了13项端到端测试用例，覆盖了从新建角色到通关的完整游戏流程。每个模块都单独验证：存档能否正常创建和恢复、6个职业的属性是否按设计赋值、商店买卖是否正确扣款和发放物品、地窟房间是否正确生成、战斗伤害计算是否准确等等。

最后是资源管理。我和组员们用AI绘画工具生成了72张高质量游戏素材，包括14张场景背景、29张角色精灵、23张敌人模型和6张角色概念图。所有素材统一规范命名并放入assets目录，还制作了免安装启动包——下载整个仓库后双击CampusRPG_GUI.exe即可直接运行，不需要安装任何Qt开发环境。

以上就是我的答辩内容，谢谢各位老师。""",

        "qa": [
            ("为什么选择暗色主题？",
             "暗色主题有三个优势：第一，游戏类应用使用暗色背景可以营造沉浸感和氛围感，契合RPG冒险的主题调性；第二，暗色背景上文字对比度更高，长时间游戏不易视觉疲劳；第三，暗色主题在近年UI设计中是主流趋势，体现了一定的设计审美。"),
            ("字体自适应缩放是如何实现的？",
             "updateUiScale()获取当前窗口尺寸，与基准尺寸1280x820比较，计算缩放比例scale=min(width/1280, height/820)。然后用scale乘以基础字号得到实际字号。所有QLabel、QPushButton、QListWidget、QTableWidget等控件的字体都按这个比例调整，保证不同分辨率下比例协调。"),
            ("排课系统中补习班的'效果减半'是什么意思？",
             "补习班允许学生在周末安排额外的课程学习。但既然是补习而非正式课程，效果只有正常的一半——即effectScale=0.5。比如专业课正常提供+2物攻+2防御的成长，补习班只提供+1物攻+1防御。这样设计是为了平衡，防止学生光靠补习班就快速提升属性而失去排课策略性。"),
            ("72张素材是怎么生成的？",
             "我们使用AI绘画工具，通过精心设计的prompt（提示词）生成统一风格的校园题材角色和场景。比如生成冰法师时描述'冰魔法师，校园风，动漫风格，蓝色调衣袍，手持法杖，透明背景'。场景图用对应的校园场景描述词。所有素材经过筛选和裁剪后统一编号命名放入assets目录。"),
            ("集成测试中最难发现的Bug是什么？",
             "最隐蔽的Bug是'轮回保留属性但装备丢失'的问题。团灭轮回后角色的经验等级和基础属性应该保留，但装备应该清空。最初版本中deserializeCharacters读取存档时，装备数据也被错误保留了。通过单独追踪equipment字段的序列化逻辑，在轮回重置时显式调用equipment.clear()解决了这个问题。"),
            ("如果窗口缩小到很小字体看不清怎么办？",
             "updateUiScale()中设置了最小缩放比例0.6，保证字体不会小于基准的60%。同时QTabWidget和QScrollArea会自动出现滚动条，用户可以用滚轮查看被遮挡的内容。实际使用中1280x820的基准尺寸已经覆盖了大多数笔记本和台式机屏幕。"),
        ],
        "images_title": "场景建模与UI界面展示",
        "images_desc": "以下为负责模块相关的7个校园场景设计模型（AI生成概念图）",
        "images": [
            "scenes/体育器材室.jpg",
            "scenes/室内体育馆.jpg",
            "scenes/废弃教室.jpg",
            "scenes/教学楼教室黄昏.jpg",
            "scenes/校园图书馆.jpg",
            "scenes/校园操场黄昏.jpg",
            "scenes/校园教学楼走廊.jpg",
        ],
        "img_width": 1.6,
        "img_per_row": 3,
    },
]

for m in members:
    doc = Document()
    title_page(doc, m["name"], m["role"], m["pct"])

    # 1. Basic Info
    H(doc, f"一、基本信息", level=1)
    add_table(doc,
        ["项目", "内容"],
        [
            ["姓名", m["name"]],
            ["角色", m["role"]],
            ["贡献度", m["pct"]],
            ["负责模块", m["modules"]],
            ["项目名称", "校园RPG冒险游戏"],
            ["技术栈", "C++17 + Qt 5.15.2 Widgets + MinGW 8.1.0"],
            ["代码规模", "C++约2,340行 + 72张素材 + INI存档系统"],
        ])
    P(doc, "")

    # 2. Detailed Work
    H(doc, f"二、工作内容详述", level=1)
    add_table(doc, ["工作内容", "具体产出与关键技术"], m["work_table"])
    P(doc, "")

    # 2.5 Image Gallery
    if "images_title" in m and "images" in m:
        H(doc, m.get("images_title", "相关模型展示"), level=2)
        P(doc, m.get("images_desc", ""), size=10, color=(100,100,100))
        P(doc, "")
        img_paths = [ASSETS + p for p in m["images"]]
        IMGS(doc, img_paths, width=m.get("img_width", 1.8), per_row=m.get("img_per_row", 3))
        P(doc, "")

    # 3. Speech Script
    H(doc, f"三、答辩发言稿", level=1)
    P(doc, "（以下为建议发言内容，可根据实际情况调整，预计用时3-5分钟）", size=10, color=(120,120,120))
    P(doc, "")
    for para in m["speech"].strip().split("\n"):
        para = para.strip()
        if para:
            P(doc, para, indent=0.7)

    doc.add_page_break()

    # 4. Q&A
    H(doc, f"四、答辩可能提问及建议回答", level=1)
    P(doc, "（以下为常见技术问题及建议回答思路，仅供参考）", size=10, color=(120,120,120))
    P(doc, "")
    for i, (q, a) in enumerate(m["qa"], 1):
        H(doc, f"Q{i}：{q}", level=3)
        P(doc, f"建议回答：{a}", indent=0.7)
        P(doc, "")

    filename = f"{m['name']}_答辩.docx"
    doc.save(BASE + filename)
    print(f"Saved: {filename}")

print("Done! 4 docs with speech scripts and Q&A sections.")
