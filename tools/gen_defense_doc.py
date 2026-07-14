import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = "Microsoft YaHei"
OUTPUT = "h:/CampusRPG_作业版/答辩分工文档.docx"

def F(run, size=12, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        s = {0:26,1:18,2:15,3:13}.get(level,13)
        F(r, size=s, bold=True)

def P(doc, text, size=12, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        F(run, size=size, bold=bold)

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: F(r, size=11, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: F(r, size=10)
    return tbl

doc = Document()

# ===== TITLE PAGE =====
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("校园RPG冒险游戏")
F(r, size=36, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("C++课程设计答辩 - 任务分工与贡献度说明")
F(r, size=18, bold=True, color=(80,80,80))

for _ in range(3): doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("计科2503班  -  2026年7月")
F(r, size=14, color=(100,100,100))

doc.add_page_break()

# ===== 1. PROJECT OVERVIEW =====
H(doc, "一、项目概述", level=1)
P(doc, "本项目是基于 C++ Qt Widgets 开发的校园题材轮回 Roguelike 回合制 RPG 游戏。玩家创建角色后经历排课学习、寒假休整、天使商店购物、七层地窟探索，团灭后进入轮回保留成长属性，最终击败伪典校长-零通关。", indent=0.7)

P(doc, "")
P(doc, "开发环境：C++17 | Qt 5.15.2 Widgets | qmake + MinGW 8.1.0 | Git + GitHub", bold=True)
P(doc, "代码规模：C++ 约2,340行 | 素材72张 | 14项完整功能 | 免安装启动包可运行", bold=True)

# ===== 2. MEMBERS =====
H(doc, "二、小组成员与贡献度", level=1)

add_table(doc,
    ["姓名", "角色", "贡献度", "核心负责模块"],
    [
        ["谢文浩", "组长", "30%", "系统架构、战斗引擎、地窟生成、等级成长、Git管理"],
        ["黄金弘", "组员", "24%", "天使/恶魔双商店、任务系统、数据持久化（INI存档序列化）"],
        ["张轩",   "组员", "23%", "角色管理、背包系统、装备穿戴、编队系统"],
        ["张申桥", "组员", "23%", "Qt图形界面（7Tab）、排课系统、集成测试、文档与资源管理"],
    ])

P(doc, "")
add_table(doc,
    ["成员", "贡献度"],
    [
        ["谢文浩(组长)", "30%"],
        ["黄金弘",       "24%"],
        ["张轩",         "23%"],
        ["张申桥",       "23%"],
    ])
P(doc, "贡献度评定依据：以代码产出量、模块复杂度、技术难点、集成工作量综合评定。组长额外承担架构设计和Git合并职责。", indent=0.7)

doc.add_page_break()

# ===== 3. DETAILED DIVISION =====
H(doc, "三、各成员详细分工说明", level=1)

# --- 谢文浩 ---
H(doc, "3.1 谢文浩（组长）- 贡献度 30%", level=2)
P(doc, "负责模块：系统架构 + 战斗系统 + 地窟探索 + 等级成长 + Git管理", bold=True)
P(doc, "")
add_table(doc,
    ["工作内容", "具体产出与关键技术"],
    [
        ["系统架构设计",
         "定义 MainWindow.h 全部数据结构（284行），包含CharacterData/EnemyData/ItemData/TaskData/RoomData等核心结构体；设计模块间接口规范；设计Profession抽象基类到6个派生职业的继承体系"],
        ["回合制战斗引擎",
         "实现 startBattle()/fightOneRound()/enemyTurn()/endBattleIfNeeded() 完整战斗流程；先手判定算法（按队伍平均等级）；伤害公式 max(1,攻击-防御)；前后排目标智能选择；6种异常状态系统（冻结/灼烧/减速/眩晕/沉默/嘲讽）；20种技能类型x6职业=120种技能效果实现"],
        ["地窟生成系统",
         "实现 buildDungeonLayer() 7层地窟配置；makeEnemyGroup() 生成7层x3类（普通/精英/BOSS）共30+种敌人及AI行为；makeLayerEquipments() 7层装备掉落表（物攻/法攻/物防/法抗四槽位）"],
        ["等级成长系统",
         "addExp() 经验累计与等级提升；属性随等级自动按职业系数增长；技能按Lv1/5/10/15/20五档解锁机制"],
        ["Git仓库管理",
         "分支策略制定、多成员代码合并与冲突解决、GitHub Releases版本打包发布"],
    ])

P(doc, "")
P(doc, "答辩重点：", bold=True)
P(doc, "1. 战斗引擎如何用C++多态实现6种异常状态的统一管理，30+种敌人的差异化AI行为。", indent=0.7)
P(doc, "2. 地窟生成算法如何保证每层难度递增且可玩性——每层敌人属性、技能、装备都不同。", indent=0.7)
P(doc, "3. Profession抽象基类+虚函数的设计如何实现新增职业只需继承不修改其他代码。", indent=0.7)

P(doc, "")

# --- 黄金弘 ---
H(doc, "3.2 黄金弘 - 贡献度 24%", level=2)
P(doc, "负责模块：商店系统 + 任务系统 + 数据持久化 + 存档序列化", bold=True)
P(doc, "")
add_table(doc,
    ["工作内容", "具体产出与关键技术"],
    [
        ["天使商店",
         "refreshAngelShop()/buyAngelItem() 实现药品/食品/养成类共14种商品；折扣累计系统（天使折扣每次+1%，最高50%）；金币交易与库存管理"],
        ["恶魔商店",
         "refreshDemonShop()/buyDemonItem()/sellEquipmentToDemon() 消耗品/装备/恶魔之友交易；恶魔币独立货币系统；装备半价回收机制；恶魔折扣与恶魔之友加成"],
        ["任务系统",
         "refreshTasks()/acceptTask()/claimTask() 完整任务生命周期；3种条件类型（击败敌人/收集物品/击败BOSS）；checkTaskCompletion() 自动进度检测；奖励发放含金币+恶魔币+物品"],
        ["数据持久化",
         "writeGame()/loadGame() 基于INI文本格式的完整存档系统；4个存档槽位独立文件（slot_1.ini~slot_4.ini）；8个序列化函数实现：角色/物品/任务/图鉴各一对读写函数"],
    ])

P(doc, "")
P(doc, "答辩重点：", bold=True)
P(doc, "1. INI文本存档的设计优势——可读性强、可手动调试、格式简单不易出错。", indent=0.7)
P(doc, "2. 8个序列化函数的分组设计——为何角色/物品/任务/图鉴各自独立而非合并。", indent=0.7)
P(doc, "3. 任务系统自动检测机制——玩家行为如何触发任务进度更新的数据流。", indent=0.7)

P(doc, "")

# --- 张轩 ---
H(doc, "3.3 张轩 - 贡献度 23%", level=2)
P(doc, "负责模块：角色管理 + 背包系统 + 装备穿戴 + 编队系统", bold=True)
P(doc, "")
add_table(doc,
    ["工作内容", "具体产出与关键技术"],
    [
        ["角色管理",
         "createCharacter() 6职业创建流程（含职业选择与名称输入）；refreshCharacters() 角色详情展示（职业图标+属性面板+技能列表）；applyProfessionStats() 根据不同职业自动赋值基础属性（HP/MP/攻/防/魔攻/法抗/活力）"],
        ["背包系统",
         "refreshInventory() 40格背包显示与容量检测；useInventoryItem() 战斗内/外使用逻辑分离（食品仅限战斗外、药品无限制）；discardInventoryItem() 删除确认；addItem() 含容量溢出提示"],
        ["装备系统",
         "equipSelectedItem() 装备穿戴/卸下切换逻辑；4种装备槽位（物攻/法攻/物防/法抗）独立管理；装备属性百分比加成计算；inventoryText() 装备信息格式化展示"],
        ["编队系统",
         "setFormationOneFront()/setFormationTwoFront() 前排1后排2与前排2后排1双模式切换；moveRoleUp()/moveRoleDown() 角色上下移动排序；toggleSelectedRoleActive() 角色启用/禁用"],
    ])

P(doc, "")
P(doc, "答辩重点：", bold=True)
P(doc, "1. 编队系统的灵活性——两种阵型如何影响战斗体验（前排多则防御面更宽）。", indent=0.7)
P(doc, "2. 装备4槽位的穿戴逻辑和属性加成计算公式——如何保证装备切换不丢属性。", indent=0.7)
P(doc, "3. 背包战斗限制的设计意图——食品不可在战斗中使用，增加策略性。", indent=0.7)

P(doc, "")

# --- 张申桥 ---
H(doc, "3.4 张申桥 - 贡献度 23%", level=2)
P(doc, "负责模块：Qt图形界面 + 排课系统 + 集成测试 + 文档与资源管理", bold=True)
P(doc, "")
add_table(doc,
    ["工作内容", "具体产出与关键技术"],
    [
        ["Qt界面开发",
         "setupSavePage()/setupGamePage() 存档首页+游戏主页双页面架构；7个Tab页布局（轮回总览/排课/角色编队/背包/任务/天使商店/地窟/轮回游记）；暗色主题CSS样式表；updateUiScale() 随窗口大小自适应缩放字体和图标；QListWidget图标显示（图鉴/编队/地窟均含对应精灵图缩略图）"],
        ["排课系统UI",
         "rebuildScheduleTable() 7天x2半天=14格课表生成；refreshSchedulePreview() 本周预计收益实时预览（活力消耗+金币收入）；工作日9科课程+周末3种行动的下拉选择；补习班规则（效果减半）"],
        ["系统集成测试",
         "全流程端到端测试：新建角色-排课-天使商店-地窟7层-最终BOSS-通关/团灭轮回；13项测试用例覆盖存档/职业/排课/商店/地窟/战斗各模块"],
        ["文档与资源",
         "README.md 完整项目文档（功能列表/编译指南/OOP特性说明）；检查点一/二/三课程设计报告；AI生成72张素材整理与规范命名（14场景+29精灵+23敌人+6角色）；免安装启动包打包（EXE+Qt5Core/Gui/Widgets+MinGW运行时+platforms插件）"],
    ])

P(doc, "")
P(doc, "答辩重点：", bold=True)
P(doc, "1. Qt GUI的Tab架构设计——7个Tab如何与后台逻辑模块对应，数据如何在Tab间流转。", indent=0.7)
P(doc, "2. 排课系统14格课表UI的实时计算——活力消耗和金币收入如何实时反馈。", indent=0.7)
P(doc, "3. 72张AI素材的生产管线和规范化命名——场景/角色/敌人素材如何与游戏中对象对应。", indent=0.7)

doc.add_page_break()

# ===== 4. TECHNICAL SUMMARY =====
H(doc, "四、项目技术亮点总览", level=1)

add_table(doc,
    ["技术点", "实现方式", "涉及成员"],
    [
        ["面向对象继承+多态", "Profession抽象基类->6派生职业，虚函数skills()/baseStats()实现运行时多态", "谢文浩"],
        ["回合制战斗引擎", "3v3前后排、先手判定、6种异常状态、120种技能效果、30+种敌人AI", "谢文浩"],
        ["INI文本存档序列化", "8个序列化函数，角色/背包/任务/图鉴完整读写，4槽位独立管理", "黄金弘"],
        ["双商店+双货币经济系统", "天使商店(金币)+恶魔商店(恶魔币)，折扣累计，装备半价回收", "黄金弘"],
        ["装备穿戴与编队系统", "4槽位装备属性百分比加成，双阵型灵活切换，角色排序", "张轩"],
        ["Qt Tab架构+暗色主题", "7个Tab页、CSS样式表、字体自适应缩放、精灵图图标展示", "张申桥"],
        ["免安装启动包", "EXE + Qt DLLs + MinGW运行时 + platforms插件，下载即玩", "张申桥"],
        ["AI素材生产管线", "AI绘画工具生成72张场景/角色/敌人素材，统一规范命名", "张申桥"],
    ])

P(doc, "")

# ===== 5. CODE STATS =====
H(doc, "五、代码与资源统计", level=1)

add_table(doc,
    ["文件", "行数/数量", "主要内容", "主要贡献者"],
    [
        ["src/MainWindow.cpp", "1,866行", "Qt界面、战斗引擎、地窟、商店、任务、背包、存档", "全员协作"],
        ["src/MainWindow.h",   "284行",   "全部数据结构定义、枚举、接口声明", "谢文浩"],
        ["include/Profession.h", "93行",  "职业抽象基类与技能结构体定义", "谢文浩"],
        ["src/Profession.cpp",  "85行",   "6个职业派生类具体实现", "谢文浩"],
        ["src/main.cpp",        "12行",   "Qt程序入口", "张申桥"],
        ["tools/build_assets.py","126行", "AI素材构建脚本", "张申桥"],
        ["README.md",           "80行",   "项目完整说明文档", "张申桥"],
        ["C++ 总计",           "2,340行", "", ""],
        ["游戏素材（assets/）", "72张",   "场景14+精灵29+敌人23+角色6", "张申桥(整理)"],
        ["启动包（release/）",  "~30MB",  "EXE+DLL+platforms+saves", "张申桥(打包)"],
    ])

doc.add_page_break()

# ===== 6. DEFENSE TIPS =====
H(doc, "六、各成员答辩准备建议", level=1)

tips = [
    ("谢文浩（组长）",
     [
         "重点展示 MainWindow.h 的数据结构设计思路——为什么选择struct+QVector而非传统的class继承",
         "战斗引擎演示：从选择敌人-进入战斗-先手判定-伤害计算-异常状态-胜利/失败的全流程",
         "地窟生成算法：如何用代码生成7层递增难度的房间和敌人配置",
         "说明职业多态设计的优势——新增职业只需继承Profession实现skills()即可，无需修改其他代码",
     ]),
    ("黄金弘",
     [
         "展示INI存档文件的可读性——打开slot_1.ini直接能看到角色数据，方便调试",
         "说明8个序列化函数的设计：为什么分成4组（角色/物品/任务/图鉴）而非一个统一函数",
         "任务系统的自动检测机制：如何在玩家击败敌人/获得物品后自动更新任务进度",
         "双商店双货币系统的设计理念——天使商店(金币)和恶魔商店(恶魔币)的经济平衡设计",
     ]),
    ("张轩",
     [
         "演示编队系统的两种阵型切换和角色排序的交互效果",
         "装备穿戴逻辑：4槽位的检查、替换、卸下流程，以及百分比加成的计算公式",
         "背包系统的使用限制设计——食品为何不能在战斗中使用（增加策略性的设计意图）",
         "说明角色创建流程中applyProfessionStats()如何根据不同职业差异化赋值属性",
     ]),
    ("张申桥",
     [
         "展示Qt界面的7个Tab页布局和暗色主题视觉效果截图",
         "排课系统14格课表UI的实时金币/活力预览计算逻辑",
         "演示免安装启动包：下载-解压-双击exe-即玩的完整体验",
         "说明AI素材生产流程：如何用AI工具高效生成72张统一风格的游戏素材",
     ]),
]

for name, items in tips:
    H(doc, name, level=2)
    for item in items:
        P(doc, "  - " + item, indent=0.7)
    P(doc, "")

# ===== SAVE =====
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
